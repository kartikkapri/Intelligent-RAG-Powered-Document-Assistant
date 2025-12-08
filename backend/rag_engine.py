"""
Advanced RAG Pipeline with Complete Document Processing
Features:
- Document ingestion: PDF, DOCX, TXT, Markdown, webpages, code files
- Clean text extraction with OCR support
- Noise removal (headers, footers, repeated sections)
- Variable chunk sizes (300-1200 tokens) based on content density
- Semantic embeddings with nomic-embed-text or mxbai-embed-large
- Chroma vector database for storage
- Context retrieval with similarity search
- Citation and source reference support
"""

import warnings
# Suppress FutureWarning about resume_download deprecation in huggingface_hub
# This is fixed in newer versions of sentence-transformers and transformers
warnings.filterwarnings('ignore', category=FutureWarning, message='.*resume_download.*')

import os
# Disable TensorFlow in transformers (we use PyTorch only)
# This prevents transformers from trying to import TensorFlow modules
# Must be set before importing transformers or sentence_transformers
os.environ.setdefault('USE_TF', '0')
os.environ.setdefault('USE_TORCH', '1')

import requests
import json
from typing import List, Dict, Tuple, Optional, Any
import numpy as np
from pathlib import Path
import io
import uuid
from database import get_database
from vector_store import QdrantVectorStore, QDRANT_AVAILABLE
import re
from datetime import datetime
import asyncio
from concurrent.futures import ThreadPoolExecutor
import hashlib

# LangChain imports
from langchain_community.document_loaders import (
    PyPDFLoader,
    UnstructuredWordDocumentLoader,
    UnstructuredExcelLoader,
    CSVLoader,
    TextLoader,
    UnstructuredMarkdownLoader,
    UnstructuredHTMLLoader,
    BSHTMLLoader,
)
from langchain.text_splitter import (
    RecursiveCharacterTextSplitter,
    MarkdownHeaderTextSplitter,
    Language,
    RecursiveCharacterTextSplitter,
)
from langchain_core.documents import Document
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_core.embeddings import Embeddings

# Workaround for onnxruntime executable stack issue on some Linux systems
# Set environment variable to allow executable stacks before importing chromadb
import os
if 'LD_PRELOAD' not in os.environ:
    # Try to work around the executable stack issue
    # This is a workaround for the "cannot enable executable stack" error
    pass

# Lazy import Chroma to avoid onnxruntime issues at import time
# We'll import it only when needed in the __init__ method
Chroma = None
CHROMA_AVAILABLE = False
try:
    # Try to import chromadb with a workaround for the executable stack issue
    # The issue occurs when chromadb tries to use onnxruntime's default embedding function
    # We'll handle this by importing lazily and providing our own embedding function
    from langchain_chroma import Chroma
    CHROMA_AVAILABLE = True
except Exception as e:
    # Chroma will be imported lazily when needed in __init__
    # This allows us to set up the environment properly before importing
    pass

# Document processing imports
import PyPDF2
from docx import Document as DocxDocument
import pandas as pd
import pdfplumber
try:
    import fitz  # PyMuPDF (imported as fitz)
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    print("⚠️ PyMuPDF (pymupdf) not available, PDF extraction may be limited")
from pdf2image import convert_from_bytes
from PIL import Image
import cv2

# Web scraping
from bs4 import BeautifulSoup

# Code processing
try:
    from pygments.lexers import get_lexer_by_name, guess_lexer_for_filename
    from pygments.util import ClassNotFound
    PYGMENTS_AVAILABLE = True
except ImportError:
    PYGMENTS_AVAILABLE = False

# OCR - Lazy import to avoid TensorFlow import at module level
# easyocr imports TensorFlow, which causes issues with transformers
EASYOCR_AVAILABLE = False
easyocr = None
# We'll import easyocr lazily only when actually needed for OCR

try:
    import pytesseract
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False

# Token counting
try:
    import tiktoken
    TIKTOKEN_AVAILABLE = True
except ImportError:
    TIKTOKEN_AVAILABLE = False
    print("⚠️ tiktoken not available, using character-based chunking")

# Embeddings
from sentence_transformers import SentenceTransformer
try:
    from sentence_transformers import CrossEncoder
except ImportError:
    CrossEncoder = None

# Hybrid search
try:
    from rank_bm25 import BM25Okapi
    BM25_AVAILABLE = True
except ImportError:
    BM25_AVAILABLE = False
    print("⚠️ rank-bm25 not available, keyword search disabled")

# NLTK for sentence tokenization
try:
    import nltk
    nltk.download('punkt', quiet=True)
    from nltk.tokenize import sent_tokenize, word_tokenize
    NLTK_AVAILABLE = True
except ImportError:
    NLTK_AVAILABLE = False
    print("⚠️ NLTK not available, using simple sentence splitting")

# Unstructured
try:
    from unstructured.partition.auto import partition
    UNSTRUCTURED_AVAILABLE = True
except ImportError:
    UNSTRUCTURED_AVAILABLE = False


class TokenCounter:
    """Token counter for chunk size management"""
    
    def __init__(self, model_name: str = "gpt-4"):
        self.encoding = None
        if TIKTOKEN_AVAILABLE:
            try:
                self.encoding = tiktoken.encoding_for_model(model_name)
            except:
                try:
                    self.encoding = tiktoken.get_encoding("cl100k_base")
                except:
                    self.encoding = None
    
    def count_tokens(self, text: str) -> int:
        """Count tokens in text"""
        if self.encoding:
            return len(self.encoding.encode(text))
        # Fallback: approximate 1 token = 4 characters
        return len(text) // 4


class AdvancedOCREngine:
    """Advanced OCR engine using EasyOCR with Tesseract fallback"""
    
    def __init__(self):
        self.easyocr_reader = None
        self.use_easyocr = False
        self.easyocr_gpu = False  # Track if GPU is enabled
        
        # Lazy import easyocr to avoid TensorFlow import at module level
        # This prevents TensorFlow from being imported before we can disable it in transformers
        try:
            import easyocr
            self.easyocr_module = easyocr
            self.use_easyocr = True
        except ImportError:
            self.easyocr_module = None
            self.use_easyocr = False
        
        # Check if GPU should be forced to CPU for OCR
        force_cpu_ocr = os.getenv('FORCE_CPU_OCR', 'false').lower() == 'true'
        
        if self.use_easyocr and self.easyocr_module:
            try:
                # Try GPU first if not forced to CPU
                if not force_cpu_ocr:
                    try:
                        import torch
                        if torch.cuda.is_available():
                            # Clear GPU cache before initializing
                            torch.cuda.empty_cache()
                            print("🔄 Initializing EasyOCR with GPU...")
                            self.easyocr_reader = self.easyocr_module.Reader(['en'], gpu=True, verbose=False)
                            self.easyocr_gpu = True
                            print("✅ EasyOCR initialized with GPU")
                        else:
                            raise Exception("CUDA not available")
                    except Exception as gpu_error:
                        # GPU initialization failed, try CPU
                        print(f"⚠️ GPU initialization failed: {gpu_error}, trying CPU...")
                        self.easyocr_reader = self.easyocr_module.Reader(['en'], gpu=False, verbose=False)
                        self.easyocr_gpu = False
                        print("✅ EasyOCR initialized with CPU")
                else:
                    # Force CPU mode
                    print("🔄 Initializing EasyOCR with CPU (forced)...")
                    self.easyocr_reader = self.easyocr_module.Reader(['en'], gpu=False, verbose=False)
                    self.easyocr_gpu = False
                    print("✅ EasyOCR initialized with CPU")
            except Exception as e:
                print(f"⚠️ EasyOCR failed: {e}, using Tesseract")
                self.use_easyocr = False
                self.easyocr_gpu = False
    
    def preprocess_image(self, image: Image.Image) -> Image.Image:
        """Preprocess image for better OCR"""
        if image.mode != 'RGB':
            image = image.convert('RGB')
        
        img_array = np.array(image)
        
        if len(img_array.shape) == 3:
            gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
        else:
            gray = img_array
        
        denoised = cv2.fastNlMeansDenoising(gray, None, 10, 7, 21)
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        enhanced = clahe.apply(denoised)
        _, binary = cv2.threshold(enhanced, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        
        return Image.fromarray(binary)
    
    def extract_text(self, image: Image.Image, filename: str = "") -> str:
        """Extract text from image with improved structure preservation"""
        try:
            processed_image = self.preprocess_image(image)
            
            if self.use_easyocr and self.easyocr_reader:
                try:
                    # Clear GPU cache before OCR if using GPU
                    if self.easyocr_gpu:
                        import torch
                        if torch.cuda.is_available():
                            torch.cuda.empty_cache()
                    
                    results = self.easyocr_reader.readtext(np.array(processed_image))
                    
                    # Improved extraction with spatial awareness
                    if not results:
                        return ""
                    
                    # Sort by vertical position (top to bottom), then horizontal (left to right)
                    # This preserves reading order and document structure
                    sorted_results = sorted(results, key=lambda x: (
                        (x[0][0][1] + x[0][2][1]) / 2,  # Average Y coordinate (top to bottom)
                        (x[0][0][0] + x[0][2][0]) / 2   # Average X coordinate (left to right)
                    ))
                    
                    text_parts = []
                    current_line = []
                    last_y = None
                    line_height_threshold = 0.05 * image.height  # 5% of image height
                    
                    for (bbox, text, confidence) in sorted_results:
                        if confidence < 0.3:
                            continue
                        
                        # Calculate bounding box properties
                        y_center = (bbox[0][1] + bbox[2][1]) / 2
                        x_center = (bbox[0][0] + bbox[2][0]) / 2
                        bbox_height = abs(bbox[2][1] - bbox[0][1])
                        
                        # Group text into lines based on vertical position
                        if last_y is None or abs(y_center - last_y) > line_height_threshold:
                            # New line - flush current line
                            if current_line:
                                # Sort current line by X position and join
                                current_line.sort(key=lambda x: x[1])  # Sort by x_center
                                line_text = ' '.join([item[0] for item in current_line])
                                text_parts.append(line_text)
                                current_line = []
                            
                            last_y = y_center
                        
                        # Add to current line
                        current_line.append((text, x_center, confidence))
                    
                    # Flush remaining line
                    if current_line:
                        current_line.sort(key=lambda x: x[1])
                        line_text = ' '.join([item[0] for item in current_line])
                        text_parts.append(line_text)
                    
                    # Join with newlines to preserve structure
                    extracted_text = '\n'.join(text_parts)
                    
                    # Post-process to improve readability
                    # Remove excessive whitespace but preserve line breaks
                    extracted_text = re.sub(r'[ \t]+', ' ', extracted_text)  # Multiple spaces to single
                    extracted_text = re.sub(r'\n{3,}', '\n\n', extracted_text)  # Multiple newlines to double
                    
                    if extracted_text.strip():
                        return extracted_text
                except RuntimeError as e:
                    # Handle CUDA out of memory errors
                    if "out of memory" in str(e).lower() or "cuda" in str(e).lower():
                        print(f"⚠️ OCR GPU out of memory: {e}")
                        # Try to clear cache and retry once
                        if self.easyocr_gpu:
                            try:
                                import torch
                                if torch.cuda.is_available():
                                    torch.cuda.empty_cache()
                                    # Try one more time with cleared cache
                                    results = self.easyocr_reader.readtext(np.array(processed_image))
                                    text_parts = []
                                    for (bbox, text, confidence) in results:
                                        if confidence > 0.3:
                                            text_parts.append(text)
                                    extracted_text = '\n'.join(text_parts)
                                    if extracted_text.strip():
                                        return extracted_text
                            except:
                                pass
                        # If retry failed, fall through to Tesseract
                        print("⚠️ Falling back to Tesseract OCR due to GPU memory issue")
                    else:
                        # Re-raise if it's not a memory error
                        raise
            
            # Fallback to Tesseract
            if TESSERACT_AVAILABLE:
                from PIL import ImageEnhance, ImageFilter
                enhanced = ImageEnhance.Contrast(processed_image).enhance(1.5)
                sharpened = enhanced.filter(ImageFilter.SHARPEN)
                text = pytesseract.image_to_string(sharpened, config='--psm 6')
                if text.strip():
                    return text
            
            return ""
        except Exception as e:
            print(f"⚠️ OCR failed: {e}")
            # Suggest the PyTorch memory configuration if it's a CUDA error
            if "cuda" in str(e).lower() or "out of memory" in str(e).lower():
                print("💡 Tip: Try setting PYTORCH_CUDA_ALLOC_CONF=expandable_segments:True to reduce fragmentation")
                print("💡 Or set FORCE_CPU_OCR=true to use CPU for OCR")
            return ""
    

class DocumentCleaner:
    """Remove noise from documents (headers, footers, repeated sections)"""
    
    def __init__(self):
        self.seen_hashes = set()
    
    def remove_headers_footers(self, text: str) -> str:
        """Remove common headers and footers"""
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line_stripped = line.strip()
            
            # Skip common header/footer patterns
            if self._is_header_footer(line_stripped):
                continue
            
            # Skip page numbers
            if re.match(r'^\d+$', line_stripped) and len(line_stripped) < 4:
                continue
            
            # Skip repeated copyright notices
            if 'copyright' in line_stripped.lower() and len(line_stripped) < 100:
                continue
            
            cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)
    
    def _is_header_footer(self, line: str) -> bool:
        """Check if line is likely a header or footer"""
        if len(line) < 3:
            return False
        
        # Common header/footer patterns
        patterns = [
            r'^page \d+',
            r'^\d+$',
            r'^confidential',
            r'^draft',
            r'^internal use only',
            r'^copyright',
        ]
        
        for pattern in patterns:
            if re.match(pattern, line.lower()):
                return True
        
        return False
    
    def remove_repeated_sections(self, text: str) -> str:
        """Remove repeated sections"""
        paragraphs = text.split('\n\n')
        seen_hashes = set()
        unique_paragraphs = []
        
        for para in paragraphs:
            para_stripped = para.strip()
            if not para_stripped or len(para_stripped) < 20:
                unique_paragraphs.append(para)
                continue
            
            # Create hash of first 100 chars
            para_hash = hashlib.md5(para_stripped[:100].lower().encode()).hexdigest()
            
            if para_hash not in seen_hashes:
                seen_hashes.add(para_hash)
                unique_paragraphs.append(para)
            else:
                # Check if it's a significant duplicate
                if len(para_stripped) > 200:  # Only skip longer duplicates
                    continue
                unique_paragraphs.append(para)
        
        return '\n\n'.join(unique_paragraphs)
    
    def clean(self, text: str) -> str:
        """Apply all cleaning operations"""
        # Remove headers/footers
        text = self.remove_headers_footers(text)
        
        # Remove repeated sections
        text = self.remove_repeated_sections(text)
        
        # Normalize whitespace
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        return text.strip()
    

class DocumentProcessor:
    """Process various document types"""
    
    def __init__(self, ocr_engine: AdvancedOCREngine):
        self.ocr_engine = ocr_engine
        self.cleaner = DocumentCleaner()
    
    async def extract_from_pdf(self, content: bytes, filename: str) -> List[Document]:
        """Extract from PDF with metadata"""
        documents = []
        
        # Try PyMuPDF first (best quality)
        if PYMUPDF_AVAILABLE:
            try:
                doc = fitz.open(stream=content, filetype="pdf")
                
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    
                    # Extract text
                    page_text = page.get_text("text")
                    
                    # Extract images and OCR
                    image_list = page.get_images()
                    ocr_texts = []
                    for img_index, img in enumerate(image_list):
                        try:
                            xref = img[0]
                            base_image = doc.extract_image(xref)
                            image_bytes = base_image["image"]
                            image = Image.open(io.BytesIO(image_bytes))
                            ocr_text = self.ocr_engine.extract_text(image, f"{filename}_p{page_num}_i{img_index}")
                            if ocr_text.strip():
                                ocr_texts.append(ocr_text)
                        except:
                            pass
                    
                    # Combine text with better structure preservation
                    full_text = page_text
                    if ocr_texts:
                        # Add OCR text with clear markers
                        full_text += "\n\n[Images/OCR Content]\n" + "\n".join(ocr_texts)
                    
                    if full_text.strip():
                        # Clean text but preserve structure
                        cleaned_text = self.cleaner.clean(full_text)
                        
                        # Preserve page context in metadata for better chunking
                        documents.append(Document(
                            page_content=cleaned_text,
                            metadata={
                                "source": filename,
                                "page": page_num + 1,
                                "type": "pdf",
                                "page_number": page_num + 1,
                                "total_pages": len(doc) if hasattr(doc, '__len__') else None
                            }
                        ))
                
                doc.close()
                if documents:
                    return documents
            except Exception as e:
                print(f"⚠️ PyMuPDF extraction failed: {e}, trying fallback...")
        
        # Fallback to pdfplumber
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text and text.strip():
                        cleaned_text = self.cleaner.clean(text)
                        documents.append(Document(
                            page_content=cleaned_text,
                            metadata={
                                "source": filename,
                                "page": page_num + 1,
                                "type": "pdf"
                            }
                        ))
            if documents:
                return documents
        except Exception as e:
            print(f"⚠️ pdfplumber extraction failed: {e}, trying PyPDF2...")
        
        # Final fallback to PyPDF2
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                if text and text.strip():
                    cleaned_text = self.cleaner.clean(text)
                    documents.append(Document(
                        page_content=cleaned_text,
                        metadata={
                            "source": filename,
                            "page": page_num + 1,
                            "type": "pdf"
                        }
                    ))
        except Exception as e:
            print(f"⚠️ PyPDF2 extraction failed: {e}")
        
        return documents
    
    async def extract_from_docx(self, content: bytes, filename: str) -> List[Document]:
        """Extract from DOCX"""
        documents = []
        
        try:
            doc = DocxDocument(io.BytesIO(content))
            
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table_num, table in enumerate(doc.tables):
                table_text = []
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        table_text.append(row_text)
                if table_text:
                    text_parts.append(f"\n[TABLE {table_num + 1}]\n" + '\n'.join(table_text))
            
            if text_parts:
                full_text = '\n\n'.join(text_parts)
                cleaned_text = self.cleaner.clean(full_text)
                
                documents.append(Document(
                    page_content=cleaned_text,
                    metadata={
                        "source": filename,
                        "type": "docx"
                    }
                ))
        except Exception as e:
            print(f"⚠️ DOCX extraction failed: {e}")
        
        return documents
    
    async def extract_from_markdown(self, content: bytes, filename: str) -> List[Document]:
        """Extract from Markdown"""
        documents = []
        
        try:
            text = content.decode('utf-8', errors='ignore')
            cleaned_text = self.cleaner.clean(text)
            
            # Use MarkdownHeaderTextSplitter to preserve structure
            headers_to_split_on = [
                ("#", "Header 1"),
                ("##", "Header 2"),
                ("###", "Header 3"),
            ]
            
            markdown_splitter = MarkdownHeaderTextSplitter(
                headers_to_split_on=headers_to_split_on
            )
            
            md_docs = markdown_splitter.split_text(cleaned_text)
            
            for i, doc in enumerate(md_docs):
                if doc.page_content.strip():
                    # Merge metadata
                    doc.metadata.update({
                        "source": filename,
                        "type": "markdown",
                        "section": i + 1
                    })
                    documents.append(doc)
        except Exception as e:
            print(f"⚠️ Markdown extraction failed: {e}")
            # Fallback to simple text extraction
            try:
                text = content.decode('utf-8', errors='ignore')
                cleaned_text = self.cleaner.clean(text)
                documents.append(Document(
                    page_content=cleaned_text,
                    metadata={"source": filename, "type": "markdown"}
                ))
            except:
                pass
        
        return documents
    
    async def extract_from_webpage(self, content: bytes, filename: str) -> List[Document]:
        """Extract from HTML/webpage"""
        documents = []
        
        try:
            html_content = content.decode('utf-8', errors='ignore')
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style", "meta", "link"]):
                script.decompose()
            
            # Extract text
            text = soup.get_text()
            
            # Clean up
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            cleaned_text = self.cleaner.clean(text)
            
            if cleaned_text.strip():
                documents.append(Document(
                    page_content=cleaned_text,
                    metadata={
                        "source": filename,
                        "type": "webpage",
                        "title": soup.title.string if soup.title else ""
                    }
                ))
        except Exception as e:
            print(f"⚠️ Webpage extraction failed: {e}")
        
        return documents
    
    async def extract_from_code(self, content: bytes, filename: str) -> List[Document]:
        """Extract from code files"""
        documents = []
        
        try:
            text = content.decode('utf-8', errors='ignore')
            
            # Try to detect language
            language = "text"
            if PYGMENTS_AVAILABLE:
                try:
                    lexer = guess_lexer_for_filename(filename, text)
                    language = lexer.name.lower()
                except:
                    pass
            
            # Clean code (preserve structure)
            lines = text.split('\n')
            cleaned_lines = []
            
            for line in lines:
                # Remove excessive whitespace but preserve indentation
                cleaned_line = line.rstrip()
                if cleaned_line:  # Keep non-empty lines
                    cleaned_lines.append(cleaned_line)
            
            cleaned_text = '\n'.join(cleaned_lines)
            
            if cleaned_text.strip():
                documents.append(Document(
                    page_content=cleaned_text,
                    metadata={
                        "source": filename,
                        "type": "code",
                        "language": language
                    }
                ))
        except Exception as e:
            print(f"⚠️ Code extraction failed: {e}")
        
        return documents
    
    async def extract_from_text(self, content: bytes, filename: str) -> List[Document]:
        """Extract from plain text"""
        documents = []
        
        try:
            text = content.decode('utf-8', errors='ignore')
            cleaned_text = self.cleaner.clean(text)
            
            if cleaned_text.strip():
                documents.append(Document(
                    page_content=cleaned_text,
                    metadata={
                        "source": filename,
                        "type": "text"
                    }
                ))
        except Exception as e:
            print(f"⚠️ Text extraction failed: {e}")
        
        return documents


class SemanticTextSplitter:
    """Advanced semantic text splitter that chunks by meaning, not just size"""
    
    def __init__(self, token_counter: TokenCounter, embedding_model: Optional[SentenceTransformer] = None):
        self.token_counter = token_counter
        self.embedding_model = embedding_model
        self.min_chunk_tokens = 100
        self.max_chunk_tokens = 800
        self.similarity_threshold = 0.5  # Threshold for semantic similarity
    
    def split_sentences(self, text: str) -> List[str]:
        """Split text into sentences"""
        if NLTK_AVAILABLE:
            try:
                return sent_tokenize(text)
            except:
                pass
        
        # Fallback: simple sentence splitting
        sentences = re.split(r'(?<=[.!?])\s+', text)
        return [s.strip() for s in sentences if s.strip()]
    
    def semantic_split(self, text: str, metadata: Dict = None) -> List[Dict]:
        """Split text semantically using embeddings"""
        if not text.strip():
            return []
        
        # Split into sentences
        sentences = self.split_sentences(text)
        if not sentences:
            return []
        
        # If no embedding model, fall back to adaptive splitting
        if self.embedding_model is None:
            return self._adaptive_split(text, metadata)
        
        # Group sentences into semantic chunks
        chunks = []
        current_chunk = []
        current_tokens = 0
        
        # Embed all sentences
        try:
            sentence_embeddings = self.embedding_model.encode(
                sentences,
                convert_to_numpy=True,
                normalize_embeddings=True,
                show_progress_bar=False
            )
        except:
            return self._adaptive_split(text, metadata)
        
        for i, (sentence, embedding) in enumerate(zip(sentences, sentence_embeddings)):
            sentence_tokens = self.token_counter.count_tokens(sentence)
            
            # If adding this sentence would exceed max size, start new chunk
            if current_tokens + sentence_tokens > self.max_chunk_tokens and current_chunk:
                chunks.append({
                    'text': ' '.join(current_chunk),
                    'sentences': current_chunk.copy(),
                    'tokens': current_tokens
                })
                current_chunk = [sentence]
                current_tokens = sentence_tokens
            else:
                # Check semantic similarity with previous sentence
                if current_chunk and i > 0:
                    prev_embedding = sentence_embeddings[i-1]
                    similarity = np.dot(embedding, prev_embedding)
                    
                    # If similarity is low and chunk is large enough, start new chunk
                    if similarity < self.similarity_threshold and current_tokens >= self.min_chunk_tokens:
                        chunks.append({
                            'text': ' '.join(current_chunk),
                            'sentences': current_chunk.copy(),
                            'tokens': current_tokens
                        })
                        current_chunk = [sentence]
                        current_tokens = sentence_tokens
                    else:
                        current_chunk.append(sentence)
                        current_tokens += sentence_tokens
                else:
                    current_chunk.append(sentence)
                    current_tokens += sentence_tokens
        
        # Add final chunk
        if current_chunk:
            chunks.append({
                'text': ' '.join(current_chunk),
                'sentences': current_chunk.copy(),
                'tokens': current_tokens
            })
        
        # Convert to Document format
        documents = []
        for i, chunk in enumerate(chunks):
            if chunk['tokens'] >= self.min_chunk_tokens:
                doc_metadata = metadata.copy() if metadata else {}
                doc_metadata.update({
                    "chunk_index": i,
                    "chunk_size_tokens": chunk['tokens'],
                    "total_chunks": len(chunks),
                    "split_method": "semantic"
                })
                documents.append(Document(
                    page_content=chunk['text'],
                    metadata=doc_metadata
                ))
        
        return documents
    
    def _adaptive_split(self, text: str, metadata: Dict = None) -> List[Document]:
        """Fallback adaptive splitting"""
        if not text.strip():
            return []
        
        # Determine chunk size
        density = self.token_counter.count_tokens(text) / len(text) if text else 0
        if density > 0.5:
            chunk_size = 600
        elif density > 0.3:
            chunk_size = 800
        else:
            chunk_size = 1000
        
        overlap = max(50, int(chunk_size * 0.2))
        
        # Use RecursiveCharacterTextSplitter
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=overlap,
            length_function=self.token_counter.count_tokens,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
        chunks = splitter.split_text(text)
        documents = []
        for i, chunk in enumerate(chunks):
            if chunk.strip() and self.token_counter.count_tokens(chunk) >= 50:
                doc_metadata = metadata.copy() if metadata else {}
                doc_metadata.update({
                    "chunk_index": i,
                    "chunk_size_tokens": self.token_counter.count_tokens(chunk),
                    "total_chunks": len(chunks),
                    "split_method": "adaptive"
                })
                documents.append(Document(
                    page_content=chunk,
                    metadata=doc_metadata
                ))
        
        return documents


class AdaptiveTextSplitter:
    """Adaptive text splitter with variable chunk sizes based on content density"""
    
    def __init__(self, token_counter: TokenCounter):
        self.token_counter = token_counter
    
    def calculate_content_density(self, text: str) -> float:
        """Calculate content density (tokens per character)"""
        tokens = self.token_counter.count_tokens(text)
        chars = len(text)
        if chars == 0:
            return 0.0
        return tokens / chars
    
    def determine_chunk_size(self, text: str, base_size: int = 800) -> int:
        """Determine chunk size based on content density"""
        density = self.calculate_content_density(text)
        
        # High density (code, technical) -> smaller chunks (300-600)
        # Medium density (normal text) -> medium chunks (600-1000)
        # Low density (sparse, lists) -> larger chunks (1000-1200)
        
        if density > 0.5:  # High density (code, technical)
            return min(600, max(300, int(base_size * 0.6)))
        elif density > 0.3:  # Medium density
            return min(1000, max(600, int(base_size * 0.8)))
        else:  # Low density
            return min(1200, max(1000, int(base_size * 1.2)))
    
    def split_text(self, text: str, metadata: Dict = None) -> List[Document]:
        """Split text with adaptive chunk sizes"""
        if not text.strip():
            return []
        
        # Determine chunk size
        chunk_size = self.determine_chunk_size(text)
        overlap = max(50, int(chunk_size * 0.2))  # 20% overlap
        
        # Use RecursiveCharacterTextSplitter
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=overlap,
            length_function=self.token_counter.count_tokens,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
        # Split
        chunks = splitter.split_text(text)
        
        # Create documents with metadata
        documents = []
        for i, chunk in enumerate(chunks):
            if chunk.strip() and self.token_counter.count_tokens(chunk) >= 50:
                doc_metadata = metadata.copy() if metadata else {}
                doc_metadata.update({
                    "chunk_index": i,
                    "chunk_size_tokens": self.token_counter.count_tokens(chunk),
                    "total_chunks": len(chunks)
                })
                documents.append(Document(
                    page_content=chunk,
                    metadata=doc_metadata
                ))
        
        return documents


class RAGEngine:
    """Advanced RAG Engine with complete pipeline"""
    
    def __init__(
        self,
        model_name: str = "gemma3:4b",
        ollama_url: str = "http://localhost:11434",
        device: Optional[str] = None,
        embedding_model: str = "nomic-ai/nomic-embed-text-v1"  # or "BAAI/bge-large-en-v1.5"
    ):
        self.model_name = model_name
        self.ollama_url = ollama_url
        self.embedding_model_name = embedding_model
        
        # Determine device
        import torch
        if os.getenv('FORCE_CPU', 'false').lower() == 'true':
            device = 'cpu'
        elif device is None:
            if torch.cuda.is_available():
                try:
                    torch.cuda.empty_cache()
                    device = 'cuda'
                    print(f"🖥️  Using GPU: {torch.cuda.get_device_name(0)}")
                except:
                    device = 'cpu'
            else:
                device = 'cpu'
        
        self.device = device
        
        # Initialize components
        print("🔄 Initializing RAG Engine...")
        
        # OCR engine
        self.ocr_engine = AdvancedOCREngine()
        
        # Document processor
        self.doc_processor = DocumentProcessor(self.ocr_engine)
        
        # Token counter
        self.token_counter = TokenCounter()
        
        # Semantic text splitter (will be initialized after embedding model)
        self.semantic_splitter = None
        
        # Adaptive text splitter (fallback)
        self.text_splitter = AdaptiveTextSplitter(self.token_counter)
        
        # Initialize embedding model
        print(f"🔄 Loading embedding model: {embedding_model}...")
        try:
            # Try to load specified model
            try:
                # Check if this is a nomic-ai model that requires trust_remote_code
                trust_remote = "nomic-ai" in embedding_model.lower()
                self.embeddings_model = SentenceTransformer(
                    embedding_model, 
                    device=device,
                    trust_remote_code=trust_remote
                )
                embedding_dim = self.embeddings_model.get_sentence_embedding_dimension()
                print(f"✅ Loaded {embedding_model} on {device} (dim: {embedding_dim})")
            except Exception as e1:
                # Fallback to alternatives
                print(f"⚠️ Failed to load {embedding_model}: {e1}")
                try:
                    # Try with correct HuggingFace path and trust_remote_code
                    self.embeddings_model = SentenceTransformer(
                        "nomic-ai/nomic-embed-text-v1", 
                        device=device,
                        trust_remote_code=True
                    )
                    embedding_dim = self.embeddings_model.get_sentence_embedding_dimension()
                    print(f"✅ Loaded nomic-ai/nomic-embed-text-v1 on {device} (dim: {embedding_dim})")
                except Exception as e2:
                    print(f"⚠️ Failed to load nomic-ai/nomic-embed-text-v1: {e2}")
                    # Final fallback
                    self.embeddings_model = SentenceTransformer("all-mpnet-base-v2", device=device)
                    embedding_dim = self.embeddings_model.get_sentence_embedding_dimension()
                    print(f"✅ Loaded all-mpnet-base-v2 on {device} (dim: {embedding_dim})")
        except Exception as e:
            print(f"⚠️ Error loading embedding model: {e}")
            self.embeddings_model = SentenceTransformer('all-MiniLM-L6-v2', device='cpu')
            embedding_dim = self.embeddings_model.get_sentence_embedding_dimension()
            self.device = 'cpu'
        
        # Initialize semantic splitter with embedding model
        self.semantic_splitter = SemanticTextSplitter(self.token_counter, self.embeddings_model)
        
        # Initialize Qdrant vector store (replaces Chroma/MongoDB for vectors)
        print("🔄 Initializing Qdrant vector database...")
        self.vector_store = None
        try:
            if QDRANT_AVAILABLE:
                qdrant_path = str(Path(__file__).parent / "qdrant_db")
                self.vector_store = QdrantVectorStore(
                    collection_name="woodai_documents",
                    path=qdrant_path,
                    embedding_dim=embedding_dim,
                    use_memory_fallback=True  # Fallback to in-memory if locked
                )
                if self.vector_store.use_memory:
                    print("⚠️  Note: Using in-memory Qdrant (data will not persist between restarts)")
                    print("   To use persistent storage, ensure no other instances are running")
                    print("   ⚠️  Previously indexed documents will NOT be available in in-memory mode!")
                else:
                    print("✅ Qdrant vector database initialized (persistent)")
                    # Check and display collection statistics
                    try:
                        collection_info = self.vector_store.get_collection_info()
                        if collection_info:
                            points_count = collection_info.get('points_count', 0)
                            if points_count > 0:
                                print(f"📚 Loaded {points_count} indexed document chunks from previous sessions")
                                print(f"   These documents are ready for RAG queries")
                            else:
                                print(f"📝 Collection is empty - ready for new document indexing")
                    except Exception as e:
                        print(f"⚠️ Could not retrieve collection info: {e}")
            else:
                print("⚠️ Qdrant not available. Install with: pip install qdrant-client")
                raise ImportError("Qdrant not available")
        except Exception as e:
            print(f"⚠️ Qdrant initialization failed: {e}")
            print(f"   Error details: {type(e).__name__}: {str(e)}")
            print("   The system will continue but vector search will be unavailable")
            import traceback
            traceback.print_exc()
            self.vector_store = None
        
        # Cross-encoder for re-ranking
        if CrossEncoder is not None:
            try:
                self.reranker = CrossEncoder('cross-encoder/ms-marco-MiniLM-L-6-v2')
                if self.device == 'cpu':
                    try:
                        self.reranker.model = self.reranker.model.to('cpu')
                    except:
                        pass
                print("✅ Cross-encoder loaded")
            except:
                self.reranker = None
        else:
            self.reranker = None
        
        # Create extracted_texts directory
        self.extracted_texts_dir = Path(__file__).parent / "extracted_texts"
        self.extracted_texts_dir.mkdir(exist_ok=True)
        
        # Knowledge base
        self.knowledge_base = []
        self.db = get_database()
        self._load_knowledge_base()
        
        print(f"✅ RAG Engine initialized")
    
    def _is_conversational_message(self, message: str) -> bool:
        """Detect if message is a conversational/greeting message that doesn't need document context"""
        message_lower = message.lower().strip()
        
        # Very short messages (likely greetings)
        if len(message_lower) <= 10:
            conversational_patterns = [
                'hey', 'hi', 'hello', 'hola', 'hey there', 'hi there',
                'howdy', 'greetings', 'good morning', 'good afternoon', 
                'good evening', 'gm', 'gn', 'sup', 'yo', 'what\'s up',
                'wassup', 'how are you', 'how are ya', 'how\'s it going',
                'thanks', 'thank you', 'thx', 'ty', 'ok', 'okay', 'k',
                'yes', 'no', 'yep', 'nope', 'sure', 'cool', 'nice', 'great'
            ]
            if message_lower in conversational_patterns:
                return True
        
        # Greeting patterns
        greeting_patterns = [
            r'^(hey|hi|hello|hola|howdy|greetings)',
            r'^(good (morning|afternoon|evening|day))',
            r'^(how are you|how\'?s it going|how are things)',
            r'^(what\'?s up|wassup|sup)',
            r'^(thanks|thank you|thx|ty)',
            r'^(ok|okay|sure|yep|nope|yes|no)$'
        ]
        
        for pattern in greeting_patterns:
            if re.match(pattern, message_lower):
                return True
        
        # Very casual/short responses
        if len(message_lower.split()) <= 3 and not any(word in message_lower for word in ['what', 'who', 'when', 'where', 'why', 'how', 'which', 'explain', 'tell', 'describe', 'what is', 'what are']):
            return True
        
        return False
    
    def _needs_document_context(self, message: str, filter_doc_ids: Optional[List[str]] = None) -> bool:
        """Determine if message needs document context retrieval"""
        # If documents are explicitly selected, always try to use them
        if filter_doc_ids and len(filter_doc_ids) > 0:
            return True
        
        # If it's a conversational message, don't search documents
        if self._is_conversational_message(message):
            return False
        
        # Check for question words or document-related queries
        message_lower = message.lower()
        question_indicators = [
            'what', 'who', 'when', 'where', 'why', 'how', 'which',
            'explain', 'describe', 'tell me', 'what is', 'what are',
            'summarize', 'analyze', 'find', 'search', 'look for',
            'according to', 'in the document', 'from the doc',
            'based on', 'document says', 'it says', 'mentioned'
        ]
        
        # If contains question indicators, likely needs context
        if any(indicator in message_lower for indicator in question_indicators):
            return True
        
        # If message is longer and seems like a substantive query
        if len(message.split()) > 5:
            return True
        
        # Default: don't search for very short messages
        return False
    
    def _is_conversational_message(self, message: str) -> bool:
        """Detect if message is a conversational/greeting message that doesn't need document context"""
        message_lower = message.lower().strip()
        
        # Very short messages (likely greetings)
        if len(message_lower) <= 10:
            conversational_patterns = [
                'hey', 'hi', 'hello', 'hola', 'hey there', 'hi there',
                'howdy', 'greetings', 'good morning', 'good afternoon', 
                'good evening', 'gm', 'gn', 'sup', 'yo', 'what\'s up',
                'wassup', 'how are you', 'how are ya', 'how\'s it going',
                'thanks', 'thank you', 'thx', 'ty', 'ok', 'okay', 'k',
                'yes', 'no', 'yep', 'nope', 'sure', 'cool', 'nice', 'great'
            ]
            if message_lower in conversational_patterns:
                return True
        
        # Greeting patterns
        greeting_patterns = [
            r'^(hey|hi|hello|hola|howdy|greetings)',
            r'^(good (morning|afternoon|evening|day))',
            r'^(how are you|how\'?s it going|how are things)',
            r'^(what\'?s up|wassup|sup)',
            r'^(thanks|thank you|thx|ty)',
            r'^(ok|okay|sure|yep|nope|yes|no)$'
        ]
        
        for pattern in greeting_patterns:
            if re.match(pattern, message_lower):
                return True
        
        # Very casual/short responses
        if len(message_lower.split()) <= 3 and not any(word in message_lower for word in ['what', 'who', 'when', 'where', 'why', 'how', 'which', 'explain', 'tell', 'describe', 'what is', 'what are']):
            return True
        
        return False
    
    def _needs_document_context(self, message: str, filter_doc_ids: Optional[List[str]] = None) -> bool:
        """Determine if message needs document context retrieval"""
        # If documents are explicitly selected, always try to use them
        if filter_doc_ids and len(filter_doc_ids) > 0:
            return True
        
        # If it's a conversational message, don't search documents
        if self._is_conversational_message(message):
            return False
        
        # Check for question words or document-related queries
        message_lower = message.lower()
        question_indicators = [
            'what', 'who', 'when', 'where', 'why', 'how', 'which',
            'explain', 'describe', 'tell me', 'what is', 'what are',
            'summarize', 'analyze', 'find', 'search', 'look for',
            'according to', 'in the document', 'from the doc',
            'based on', 'document says', 'it says', 'mentioned'
        ]
        
        # If contains question indicators, likely needs context
        if any(indicator in message_lower for indicator in question_indicators):
            return True
        
        # If message is longer and seems like a substantive query
        if len(message.split()) > 5:
            return True
        
        # Default: don't search for very short messages
        return False
    
    def _load_knowledge_base(self):
        """Load documents from knowledge base directory"""
        kb_path = Path("knowledge_base")
        
        if not kb_path.exists():
            kb_path.mkdir(parents=True)
            sample_doc = kb_path / "sample.txt"
            sample_doc.write_text("Welcome to WoodAI!")
    
    async def _extract_documents(self, content: bytes, filename: str) -> List[Document]:
        """Extract documents based on file type"""
        file_ext = filename.lower().split('.')[-1] if '.' in filename else ''
        
        # PDF
        if file_ext == 'pdf':
            return await self.doc_processor.extract_from_pdf(content, filename)
        
        # DOCX
        elif file_ext in ['docx', 'doc']:
            return await self.doc_processor.extract_from_docx(content, filename)
        
        # Markdown
        elif file_ext in ['md', 'markdown']:
            return await self.doc_processor.extract_from_markdown(content, filename)
        
        # HTML/Webpage
        elif file_ext in ['html', 'htm']:
            return await self.doc_processor.extract_from_webpage(content, filename)
        
        # Code files
        elif file_ext in ['py', 'js', 'ts', 'java', 'cpp', 'c', 'go', 'rs', 'rb', 'php', 'swift', 'kt', 'scala', 'r', 'sql', 'sh', 'bash', 'yaml', 'yml', 'json', 'xml', 'css', 'html', 'vue', 'jsx', 'tsx']:
            return await self.doc_processor.extract_from_code(content, filename)
        
        # Text
        elif file_ext == 'txt':
            return await self.doc_processor.extract_from_text(content, filename)
        
        # Images (OCR)
        elif file_ext in ['jpg', 'jpeg', 'png', 'bmp', 'tiff', 'tif', 'gif']:
            try:
                image = Image.open(io.BytesIO(content))
                text = self.ocr_engine.extract_text(image, filename)
                if text.strip():
                    cleaned = self.doc_processor.cleaner.clean(text)
                    return [Document(
                        page_content=cleaned,
                        metadata={"source": filename, "type": "image"}
                    )]
            except:
                pass
        
        # Default: try as text
        return await self.doc_processor.extract_from_text(content, filename)
    
    async def index_document(self, content: bytes, filename: str) -> Dict:
        """Index a document with full pipeline"""
        import torch
        
        try:
            if torch.cuda.is_available():
                torch.cuda.empty_cache()
            
            print(f"📄 Processing document: {filename}")
            print(f"  File size: {len(content)} bytes")
            
            # Step 1: Extract documents
            documents = await self._extract_documents(content, filename)
            
            if not documents:
                return {
                    'success': False,
                    'error': f'Could not extract content from {filename}'
                }
            
            print(f"✅ Extracted {len(documents)} document sections")
            
            # Step 2: Chunk documents with semantic splitting (preferred) or adaptive
            all_chunks = []
            for doc in documents:
                # Try semantic splitting first
                if self.semantic_splitter:
                    try:
                        chunks = self.semantic_splitter.semantic_split(doc.page_content, doc.metadata)
                        if chunks:
                            all_chunks.extend(chunks)
                            continue
                    except Exception as e:
                        print(f"⚠️ Semantic splitting failed, using adaptive: {e}")
                
                # Fallback to adaptive splitting
                chunks = self.text_splitter.split_text(doc.page_content, doc.metadata)
                all_chunks.extend(chunks)
            
            if not all_chunks:
                return {
                    'success': False,
                    'error': 'No valid chunks created'
                }
            
            print(f"✅ Created {len(all_chunks)} chunks with adaptive sizing")
            
            # Step 3: Generate embeddings
            print(f"🔄 Generating embeddings for {len(all_chunks)} chunks...")
            chunk_texts = [chunk.page_content for chunk in all_chunks]
            
            # Use the same embedding model as initialization to ensure dimension consistency
            # Move to CPU for indexing to avoid GPU memory issues
            embedding_model = self.embeddings_model
            if hasattr(embedding_model, 'to'):
                # Temporarily move to CPU if on GPU
                original_device = next(embedding_model.parameters()).device if hasattr(embedding_model, 'parameters') else None
                try:
                    embedding_model = embedding_model.to('cpu')
                except:
                    pass  # If already on CPU or can't move, continue
            
            # Verify embedding dimension matches Qdrant collection
            if self.vector_store:
                actual_dim = embedding_model.get_sentence_embedding_dimension()
                expected_dim = self.vector_store.embedding_dim
                if actual_dim != expected_dim:
                    print(f"⚠️ Embedding dimension mismatch!")
                    print(f"   Model dimension: {actual_dim}, Qdrant collection dimension: {expected_dim}")
                    print(f"   Recreating Qdrant collection with correct dimension...")
                    # Recreate collection with correct dimension
                    try:
                        self.vector_store.client.delete_collection(self.vector_store.collection_name)
                        self.vector_store.embedding_dim = actual_dim
                        self.vector_store._ensure_collection()
                        print(f"✅ Qdrant collection recreated with dimension {actual_dim}")
                    except Exception as e:
                        print(f"⚠️ Could not recreate collection: {e}")
                        print(f"   Continuing with existing collection (may cause errors)")
            
            batch_size = 32
            all_embeddings = []
            
            for i in range(0, len(chunk_texts), batch_size):
                batch = chunk_texts[i:i + batch_size]
                try:
                    embeddings = embedding_model.encode(
                        batch,
                        batch_size=len(batch),
                        show_progress_bar=True,
                        convert_to_numpy=True,
                        normalize_embeddings=True
                    )
                    all_embeddings.append(embeddings)
                    print(f"  ✅ Processed {min(i + batch_size, len(chunk_texts))}/{len(chunk_texts)} chunks...")
                except Exception as e:
                    print(f"❌ Error in batch: {e}")
                    raise
            
            embeddings_array = np.vstack(all_embeddings) if all_embeddings else np.array([])
            
            # Verify embedding dimension before storing
            if len(embeddings_array) > 0:
                actual_dim = embeddings_array.shape[1]
                print(f"📊 Generated embeddings with dimension: {actual_dim}")
            
            # Step 4: Store in Qdrant (primary vector store)
            doc_id = str(uuid.uuid4())
            if self.vector_store:
                # Check dimension compatibility
                if len(embeddings_array) > 0:
                    actual_dim = embeddings_array.shape[1]
                    expected_dim = self.vector_store.embedding_dim
                    
                    if actual_dim != expected_dim:
                        print(f"⚠️ Embedding dimension mismatch detected!")
                        print(f"   Generated embeddings: {actual_dim}D")
                        print(f"   Qdrant collection expects: {expected_dim}D")
                        print(f"   Attempting to recreate collection with correct dimension...")
                        try:
                            # Delete and recreate collection with correct dimension
                            collection_name = self.vector_store.collection_name
                            print(f"   Deleting existing collection '{collection_name}'...")
                            self.vector_store.client.delete_collection(collection_name)
                            print(f"   Updating embedding dimension to {actual_dim}...")
                            self.vector_store.embedding_dim = actual_dim
                            print(f"   Recreating collection...")
                            self.vector_store._ensure_collection()
                            print(f"✅ Qdrant collection recreated with dimension {actual_dim}")
                        except Exception as recreate_error:
                            print(f"❌ Could not recreate collection: {recreate_error}")
                            import traceback
                            traceback.print_exc()
                            print(f"\n   Manual fix required:")
                            print(f"   1. Stop the application")
                            print(f"   2. Delete Qdrant database: rm -rf {self.vector_store.path}")
                            print(f"   3. Restart the application")
                            raise ValueError(
                                f"Embedding dimension mismatch: {actual_dim}D vs {expected_dim}D. "
                                f"Collection recreation failed: {recreate_error}. "
                                f"Please delete Qdrant database manually."
                            )
                
                print(f"💾 Storing {len(all_chunks)} chunks in Qdrant...")
                try:
                    # Prepare texts, embeddings, and metadatas
                    texts = [chunk.page_content for chunk in all_chunks]
                    metadatas = []
                    for chunk in all_chunks:
                        meta = chunk.metadata.copy()
                        meta['doc_id'] = doc_id
                        metadatas.append(meta)
                    
                    # Add to Qdrant
                    self.vector_store.add_documents(
                        texts=texts,
                        embeddings=embeddings_array,
                        metadatas=metadatas
                    )
                    print("✅ Documents stored in Qdrant")
                except Exception as e:
                    print(f"⚠️ Qdrant storage failed: {e}")
                    import traceback
                    traceback.print_exc()
            else:
                print("⚠️ Vector store not available, skipping storage")
            
            # Step 5: Store document metadata in MongoDB (for tracking, not vectors)
            try:
                await self.db.documents.insert_one({
                    "doc_id": doc_id,
                    "filename": filename,
                    "chunk_count": len(all_chunks),
                    "created_at": datetime.utcnow()
                })
                print(f"✅ Document metadata saved to MongoDB")
            except Exception as e:
                print(f"⚠️ MongoDB metadata save failed: {e}")
            
            # Save extracted text for logging
            try:
                safe_filename = re.sub(r'[^\w\.-]', '_', filename)
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                log_filepath = self.extracted_texts_dir / f"{timestamp}_{safe_filename}.txt"
                
                with open(log_filepath, 'w', encoding='utf-8') as f:
                    f.write(f"=== EXTRACTED TEXT LOG ===\n")
                    f.write(f"Filename: {filename}\n")
                    f.write(f"Date: {datetime.now().isoformat()}\n")
                    f.write(f"Chunks: {len(all_chunks)}\n")
                    f.write(f"{'='*50}\n\n")
                    for chunk in all_chunks:
                        f.write(f"\n--- Chunk {chunk.metadata.get('chunk_index', '?')} ---\n")
                        f.write(chunk.page_content)
                        f.write("\n\n")
            except Exception as e:
                print(f"⚠️ Could not save log: {e}")
            
            return {
                'success': True,
                'doc_id': doc_id,
                'filename': filename,
                'chunks': len(all_chunks),
                'sections': len(documents)
            }
        
        except Exception as e:
            print(f"❌ Error indexing document: {e}")
            import traceback
            traceback.print_exc()
            if torch.cuda.is_available():
                torch.cuda.empty_cache()
            return {
                'success': False,
                'error': str(e)
            }
    
    async def retrieve_context(self, query: str, top_k: int = 5, use_hybrid: bool = True, min_similarity: float = 0.2, filter_doc_ids: Optional[List[str]] = None) -> Tuple[str, List[Dict]]:
        """Retrieve context with citations using Qdrant hybrid search
        
        Args:
            query: Search query
            top_k: Number of results to return
            use_hybrid: Whether to use hybrid search
            min_similarity: Minimum similarity score
            filter_doc_ids: Optional list of doc_ids to filter by
        """
        try:
            # Check if vector store is available and has data
            if not self.vector_store:
                print("❌ Vector store not available - cannot retrieve context")
                return "", []
            
            # Verify collection has data
            try:
                collection_info = self.vector_store.get_collection_info()
                points_count = collection_info.get('points_count', 0) if collection_info else 0
                if points_count == 0:
                    print("⚠️ Vector store collection is empty - no documents indexed yet")
                    print("   Please upload and index documents first")
                    return "", []
                else:
                    print(f"🔍 Searching in {points_count} indexed document chunks...")
            except Exception as e:
                print(f"⚠️ Could not verify collection: {e}")
            
            # Generate query embedding
            query_embedding = self.embeddings_model.encode(
                query,
                convert_to_numpy=True,
                normalize_embeddings=True
            )
            
            results = []
            
            # Search in Qdrant (primary vector store)
            if self.vector_store:
                try:
                    # Build filter conditions if doc_ids are specified
                    filter_conditions = None
                    if filter_doc_ids and len(filter_doc_ids) > 0:
                        # Pass list of doc_ids to filter
                        filter_conditions = {"doc_id": filter_doc_ids}
                        print(f"🔍 Filtering to {len(filter_doc_ids)} document(s): {filter_doc_ids}")
                    
                    if use_hybrid and hasattr(self.vector_store, 'hybrid_search'):
                        # Use hybrid search (semantic + keyword)
                        print(f"🔍 Hybrid search: semantic + keyword matching")
                        results = self.vector_store.hybrid_search(
                            query_embedding=query_embedding,
                            query_text=query,
                            top_k=top_k * 3,  # Get more for filtering and re-ranking
                            semantic_weight=0.7,
                            keyword_weight=0.3,
                            filter_conditions=filter_conditions
                        )
                    else:
                        # Use semantic search only
                        print(f"🔍 Semantic search only")
                        results = self.vector_store.search(
                            query_embedding=query_embedding,
                            top_k=top_k * 3,
                            score_threshold=min_similarity,  # Minimum similarity threshold (lowered for more context)
                            filter_conditions=filter_conditions
                        )
                    
                    
                    # Mark source
                    for result in results:
                        result['source'] = 'qdrant'
                    
                    # Filter by doc_ids if specified (post-filter for hybrid search or if MatchAny not available)
                    if filter_doc_ids and len(filter_doc_ids) > 0:
                        original_count = len(results)
                        results = [r for r in results if r.get('doc_id') in filter_doc_ids]
                        if original_count != len(results):
                            print(f"🔍 Post-filtered to {len(results)} results from {len(filter_doc_ids)} selected document(s)")
                    
                    print(f"✅ Found {len(results)} results from Qdrant")
                except AttributeError as e:
                    error_msg = (
                        f"❌ Qdrant API Error in RAG: {e}\n"
                        f"   This indicates the Qdrant client method is not available.\n"
                        f"   Solutions:\n"
                        f"   1. Upgrade qdrant-client: pip install --upgrade qdrant-client\n"
                        f"   2. Check Qdrant client initialization\n"
                        f"   3. Verify qdrant-client version compatibility"
                    )
                    print(error_msg)
                    import traceback
                    traceback.print_exc()
                    results = []  # Fallback to empty results
                except RuntimeError as e:
                    error_msg = (
                        f"❌ Qdrant Runtime Error in RAG: {e}\n"
                        f"   This indicates a problem with Qdrant database access.\n"
                        f"   Solutions:\n"
                        f"   1. Check if Qdrant database is accessible\n"
                        f"   2. Verify collection exists\n"
                        f"   3. Try restarting the backend\n"
                        f"   4. Check database lock files"
                    )
                    print(error_msg)
                    results = []  # Fallback to empty results
                except Exception as e:
                    error_msg = (
                        f"⚠️ Qdrant search failed: {type(e).__name__}: {e}\n"
                        f"   Query: {query[:50]}...\n"
                        f"   Falling back to empty results. Check Qdrant connection."
                    )
                    print(error_msg)
                    import traceback
                    traceback.print_exc()
                    results = []  # Fallback to empty results
                    import traceback
                    traceback.print_exc()
            
            # Re-rank with cross-encoder if available (improves accuracy)
            if self.reranker and results:
                print(f"🔄 Re-ranking {len(results)} results with cross-encoder...")
                pairs = [[query, result['text']] for result in results]
                try:
                    scores = self.reranker.predict(pairs)
                    for i, result in enumerate(results):
                        # Combine Qdrant score with reranker score
                        rerank_score = float(scores[i])
                        original_score = result.get('similarity', 0.0)
                        # Weighted combination: 60% reranker, 40% original
                        result['similarity'] = 0.6 * rerank_score + 0.4 * original_score
                        result['rerank_score'] = rerank_score
                    results.sort(key=lambda x: x['similarity'], reverse=True)
                    print(f"✅ Re-ranking complete")
                except Exception as e:
                    print(f"⚠️ Re-ranking failed: {e}")
            
            # Expand context by including surrounding chunks from same document
            expanded_results = self._expand_context_with_surrounding_chunks(results, top_k)
            
            # Get top-k expanded results
            top_results = expanded_results[:top_k]
            
            print(f"📋 Processing {len(top_results)} top results (with context expansion) for context formatting")
            
            # Format context with citations - now includes expanded context
            context_parts = []
            citations = []
            seen_chunks = set()  # Avoid duplicate chunks
            
            for i, result in enumerate(top_results):
                text = result.get('text', '')
                chunk_id = result.get('chunk_id', f"{result.get('doc_id', '')}_{result.get('chunk_index', i)}")
                
                if not text or not text.strip():
                    print(f"⚠️ Skipping result {i+1}: empty text")
                    continue
                
                # Skip if we've already included this chunk
                if chunk_id in seen_chunks:
                    continue
                seen_chunks.add(chunk_id)
                    
                metadata = result.get('metadata', {})
                source = metadata.get('source', result.get('source', 'unknown'))
                page = metadata.get('page', '')
                section = metadata.get('section', '')
                doc_id = result.get('doc_id', '')
                chunk_index = result.get('chunk_index', '')
                is_expanded = result.get('is_expanded_context', False)
                
                # Create citation
                citation = f"[{i+1}]"
                if page:
                    citation += f" {source}, page {page}"
                elif section:
                    citation += f" {source}, section {section}"
                else:
                    citation += f" {source}"
                
                if is_expanded:
                    citation += " (expanded context)"
                
                citations.append({
                    'number': i + 1,
                    'text': text[:200] + "..." if len(text) > 200 else text,
                    'source': source,
                    'page': page,
                    'section': section,
                    'doc_id': doc_id,
                    'chunk_index': chunk_index,
                    'similarity': result.get('similarity', 0.0),
                    'rerank_score': result.get('rerank_score', None),
                    'is_expanded': is_expanded
                })
                
                context_parts.append(f"{citation}\n{text}")
            
            context = '\n\n'.join(context_parts)
            
            if context:
                context_length = len(context)
                print(f"✅ Context formatted: {len(citations)} chunks, {context_length} characters")
                print(f"   Context preview: {context[:200]}...")
            else:
                print(f"⚠️ No context generated from {len(top_results)} results")
            
            return context, citations
        
        except Exception as e:
            print(f"⚠️ Error retrieving context: {e}")
            return "", []
    
    def _expand_context_with_surrounding_chunks(self, results: List[Dict], top_k: int) -> List[Dict]:
        """Expand context by including surrounding chunks from the same document
        
        This helps the LLM make decisions based on the whole document context,
        not just isolated chunks. It scans through the entire document context
        to provide comprehensive information.
        """
        if not results or not self.vector_store:
            return results
        
        expanded_results = []
        processed_docs = {}  # Track which doc_id + chunk_index combinations we've processed
        
        # Group results by doc_id to find neighbors more efficiently
        doc_groups = {}
        for result in results[:top_k * 2]:  # Process more results to find neighbors
            doc_id = result.get('doc_id')
            if doc_id:
                if doc_id not in doc_groups:
                    doc_groups[doc_id] = []
                doc_groups[doc_id].append(result)
        
        # Process each document group
        for doc_id, doc_results in doc_groups.items():
            try:
                # Get all chunks from this document
                filter_conditions = {"doc_id": [doc_id]}
                
                # Use a representative query embedding (average of top results)
                query_texts = [r.get('text', '') for r in doc_results[:3] if r.get('text')]
                if query_texts:
                    query_embedding = self.embeddings_model.encode(
                        ' '.join(query_texts),
                        convert_to_numpy=True,
                        normalize_embeddings=True
                    )
                else:
                    continue
                
                # Search for all chunks from this document
                all_doc_chunks = self.vector_store.search(
                    query_embedding=query_embedding,
                    top_k=200,  # Get many chunks to find neighbors
                    filter_conditions=filter_conditions,
                    score_threshold=0.0  # Get all chunks regardless of similarity
                )
                
                # Create a map of chunk_index -> chunk for quick lookup
                chunk_map = {}
                for chunk in all_doc_chunks:
                    chunk_idx = chunk.get('chunk_index')
                    if chunk_idx is not None:
                        chunk_map[chunk_idx] = chunk
                
                # For each result, find and add neighboring chunks
                for result in doc_results:
                    chunk_index = result.get('chunk_index')
                    
                    if chunk_index is None:
                        # Can't expand without chunk_index
                        if (doc_id, -1) not in processed_docs:
                            expanded_results.append(result)
                            processed_docs[(doc_id, -1)] = True
                        continue
                    
                    key = (doc_id, chunk_index)
                    if key in processed_docs:
                        continue
                    processed_docs[key] = True
                    
                    # Add the original result
                    expanded_results.append(result)
                    
                    # Find neighboring chunks (2 before and 2 after)
                    neighbor_indices = [
                        chunk_index - 2, chunk_index - 1,  # Previous chunks
                        chunk_index + 1, chunk_index + 2   # Next chunks
                    ]
                    
                    neighbors = []
                    for neighbor_idx in neighbor_indices:
                        if neighbor_idx in chunk_map:
                            neighbor = chunk_map[neighbor_idx].copy()
                            neighbor_key = (neighbor.get('doc_id'), neighbor_idx)
                            
                            # Only add if not already processed
                            if neighbor_key not in processed_docs:
                                # Mark as expanded context
                                neighbor['is_expanded_context'] = True
                                neighbor['original_similarity'] = result.get('similarity', 0.0)
                                # Slightly lower similarity for expanded chunks
                                neighbor['similarity'] = result.get('similarity', 0.0) * 0.85
                                neighbors.append(neighbor)
                    
                    # Sort neighbors by chunk_index to maintain document order
                    neighbors.sort(key=lambda x: x.get('chunk_index', 0))
                    
                    # Add neighbors (limit to 4 total: 2 before, 2 after)
                    for neighbor in neighbors[:4]:
                        neighbor_key = (neighbor.get('doc_id'), neighbor.get('chunk_index'))
                        if neighbor_key not in processed_docs:
                            expanded_results.append(neighbor)
                            processed_docs[neighbor_key] = True
                            
            except Exception as e:
                print(f"⚠️ Error expanding context for document {doc_id}: {e}")
                # Add original results if expansion fails
                for result in doc_results:
                    key = (result.get('doc_id'), result.get('chunk_index', -1))
                    if key not in processed_docs:
                        expanded_results.append(result)
                        processed_docs[key] = True
                continue
        
        # Re-sort by similarity (original results first, then expanded)
        expanded_results.sort(key=lambda x: (
            not x.get('is_expanded_context', False),  # Original results first
            -x.get('similarity', 0.0)  # Then by similarity descending
        ))
        
        print(f"📚 Context expansion: {len(results)} original → {len(expanded_results)} expanded chunks")
        return expanded_results
    
    async def generate_response(
        self,
        message: str,
        context_length: int,
        memory_enabled: bool,
        temperature: float,
        system_prompt: str,
        history: List[Dict],
        top_k: int = 5,
        filter_doc_ids: Optional[List[str]] = None
    ):
        """Generate response with citations"""
        try:
            if not self.is_available():
                return (
                    "⚠️ Ollama is not running. Please start it with:\n\n"
                    "1. Open a terminal\n"
                    "2. Run: ollama serve\n"
                    "3. Ensure model is pulled: ollama pull gemma2:2b"
                )
            
            query = message.strip()
            if not query:
                return "Please provide a question or query."
            
            # Determine if this needs document context
            needs_context = self._needs_document_context(query, filter_doc_ids)
            context = ""
            citations = []
            
            if needs_context:
                # Retrieve context with citations
                # Increase top_k for better context (especially for summarization and analysis)
                query_lower = query.lower()
                is_analysis_query = any(word in query_lower for word in ["summar", "conclusion", "analyze", "synthesize", "overview", "key points", "main ideas", "summary"])
                # Use more chunks for analysis queries to get better context for summarization
                retrieval_k = top_k * 2 if is_analysis_query else top_k
                
                print(f"🔍 Retrieving context for query: '{query[:100]}...'")
                if filter_doc_ids:
                    print(f"📄 Filtering to selected documents: {filter_doc_ids}")
                else:
                    print(f"📚 Searching all documents in knowledge base")
                
                context, citations = await self.retrieve_context(query, top_k=retrieval_k, min_similarity=0.2, filter_doc_ids=filter_doc_ids)
                
                if context:
                    print(f"✅ Retrieved {len(citations)} relevant chunks ({len(context)} chars)")
                    print(f"   Context will be pipelined to LLM")
                else:
                    print("⚠️ No relevant context found - LLM will answer without document context")
            else:
                print(f"💬 Conversational message detected - responding naturally without document search")
            
            # Build conversation history
            conversation = []
            if memory_enabled and len(history) > 1:
                max_history = min(len(history) - 1, context_length // 512)
                for msg in history[-max_history:]:
                    role = msg.get('role', 'user')
                    content = msg.get('content', '')
                    if role == 'user':
                        conversation.append(f"User: {content}")
                    elif role == 'assistant':
                        conversation.append(f"Assistant: {content}")
            
            # Build prompt
            prompt = f"{system_prompt}\n\n"
            
            if context and needs_context:
                print(f"📝 Adding context to LLM prompt ({len(context)} characters)")
                prompt += f"=== RELEVANT DOCUMENT CONTEXT ===\n"
                prompt += f"The following information has been retrieved from {len(citations)} document chunks (including expanded context) from the user's documents.\n"
                prompt += f"**IMPORTANT: Read and analyze the ENTIRE context below before answering.**\n"
                prompt += f"The context includes both directly relevant chunks and surrounding chunks from the same documents to provide complete context.\n\n"
                prompt += f"{context}\n\n"
                prompt += f"=== ANSWERING INSTRUCTIONS ===\n"
                prompt += f"You are an intelligent assistant. Follow these steps to answer the question:\n\n"
                prompt += f"**STEP 1: Analyze the COMPLETE context FIRST**\n"
                prompt += f"1. Read through ALL the document context provided above - do not focus on just one line or chunk.\n"
                prompt += f"2. Understand the FULL context and how different chunks relate to each other.\n"
                prompt += f"3. Scan through the entire context to find all relevant information, not just the first matching line.\n"
                prompt += f"4. Consider the relationships between different parts of the document.\n"
                prompt += f"5. Make your decision based on the WHOLE context, synthesizing information from multiple chunks.\n\n"
                prompt += f"**STEP 2: Answer comprehensively from the full context**\n"
                prompt += f"6. If the context has sufficient information:\n"
                prompt += f"   - Answer based PRIMARILY on the COMPLETE context provided (not just single lines)\n"
                prompt += f"   - Cite sources using [1], [2], etc. when referencing specific information\n"
                prompt += f"   - Synthesize information from MULTIPLE chunks to provide a comprehensive answer\n"
                prompt += f"   - Draw logical conclusions from the ENTIRE provided context\n"
                prompt += f"   - If information appears in multiple places, combine and reconcile it\n\n"
                prompt += f"**STEP 3: If context is insufficient or doesn't answer the question**\n"
                prompt += f"7. If the context doesn't contain enough information to fully answer the question:\n"
                prompt += f"   - First, state what information IS available in the context (if any)\n"
                prompt += f"   - Then, clearly indicate that the answer is not fully available in the provided documents\n"
                prompt += f"   - Use this format: \"The provided documents don't contain [specific information], but I can tell you from my knowledge that...\"\n"
                prompt += f"   - Provide your answer using your general knowledge\n"
                prompt += f"   - Be helpful and comprehensive in your response\n\n"
                prompt += f"**IMPORTANT GUIDELINES:**\n"
                prompt += f"- Always prioritize context when it's available and relevant\n"
                prompt += f"- When using general knowledge, always acknowledge it clearly\n"
                prompt += f"- If context partially answers the question, use it and supplement with general knowledge if needed\n"
                prompt += f"- Be transparent about what comes from documents vs. your knowledge\n"
                prompt += f"- ALWAYS scan through the ENTIRE context, not just the first matching line\n\n"
            elif needs_context and not context:
                # Document search was attempted but no context found
                prompt += f"=== ANSWERING INSTRUCTIONS ===\n"
                prompt += f"No relevant documents were found in the knowledge base for this question.\n\n"
                prompt += f"You should:\n"
                prompt += f"1. Clearly state: \"The provided documents don't contain information about this topic.\"\n"
                prompt += f"2. Then provide a helpful answer using your general knowledge\n"
                prompt += f"3. Use this format: \"However, I can tell you from my knowledge that...\"\n"
                prompt += f"4. Be comprehensive and helpful in your response\n\n"
            else:
                # Conversational message - respond naturally
                prompt += f"=== INSTRUCTIONS ===\n"
                prompt += f"This is a conversational message. Respond naturally and helpfully without mentioning documents.\n"
                prompt += f"Just have a friendly conversation - no need to reference documents or context.\n\n"
            
            if conversation:
                prompt += f"=== CONVERSATION HISTORY ===\n" + '\n'.join(conversation) + "\n\n"
            
            prompt += f"=== QUESTION ===\n{query}\n\n=== ANSWER ===\n"
            
            # Calculate required context window size
            # Estimate tokens: ~4 characters per token, add buffer for safety
            prompt_tokens = len(prompt) // 4
            required_ctx = max(prompt_tokens + 1000, context_length)  # Ensure enough space for response
            
            print(f"📏 Prompt size: {len(prompt)} chars (~{prompt_tokens} tokens), setting context window to {required_ctx}")
            if context:
                print(f"🔍 Prompt preview (first 500 chars): {prompt[:500]}...")
                print(f"🔍 Prompt contains context: {'RELEVANT DOCUMENT CONTEXT' in prompt}")
            
            # Call Ollama with retry logic for model loading errors
            max_retries = 2
            for attempt in range(max_retries):
                try:
                    response = requests.post(
                        f"{self.ollama_url}/api/generate",
                        json={
                            "model": self.model_name,
                            "prompt": prompt,
                            "stream": False,
                            "options": {
                                "temperature": temperature,
                                "num_ctx": required_ctx,  # Context window size - CRITICAL for RAG!
                                "num_predict": min(context_length, 2048)  # Max tokens to generate
                            }
                        },
                        timeout=120
                    )
                    
                    if response.status_code == 200:
                        result = response.json()
                        answer = result.get('response', 'No response generated')
                        
                        # Append citations if available
                        if citations:
                            answer += "\n\n=== Sources ===\n"
                            for citation in citations:
                                answer += f"{citation['number']}. {citation['source']}"
                                if citation.get('page'):
                                    answer += f", page {citation['page']}"
                                answer += f" (similarity: {citation['similarity']:.3f})\n"
                        
                        return answer
                    else:
                        error_text = response.text
                        # Check for model loading errors
                        if "load" in error_text.lower() and ("EOF" in error_text or "connection" in error_text.lower()):
                            if attempt < max_retries - 1:
                                print(f"⚠️ Model loading error detected, attempting to preload model (attempt {attempt + 1}/{max_retries})...")
                                self._preload_model()
                                continue  # Retry
                            else:
                                return (
                                    f"Error: Model loading failed. This usually means:\n"
                                    f"1. The model '{self.model_name}' is not available or corrupted\n"
                                    f"2. Ollama is having trouble loading the model\n"
                                    f"3. There's insufficient memory/disk space\n\n"
                                    f"Try:\n"
                                    f"- Check if model exists: ollama list\n"
                                    f"- Restart Ollama: ollama serve\n"
                                    f"- Re-pull the model: ollama pull {self.model_name}\n"
                                    f"Original error: {response.status_code} - {error_text}"
                                )
                        return f"Error: {response.status_code} - {error_text}"
                        
                except requests.exceptions.ConnectionError as e:
                    if attempt < max_retries - 1:
                        print(f"⚠️ Connection error, retrying (attempt {attempt + 1}/{max_retries})...")
                        import time
                        time.sleep(1)
                        continue
                    else:
                        raise
            
        except requests.exceptions.ConnectionError as e:
            error_msg = (
                f"❌ Cannot connect to Ollama.\n\n"
                f"Please ensure Ollama is running: ollama serve\n"
                f"Error: {str(e)}"
            )
            return error_msg
        except Exception as e:
            error_str = str(e)
            # Check for EOF/model loading errors in exception message
            if "load" in error_str.lower() and ("EOF" in error_str or "connection" in error_str.lower()):
                return (
                    f"Error: Model loading failed. This usually means:\n"
                    f"1. The model '{self.model_name}' is not available or corrupted\n"
                    f"2. Ollama is having trouble loading the model\n"
                    f"3. There's insufficient memory/disk space\n\n"
                    f"Try:\n"
                    f"- Check if model exists: ollama list\n"
                    f"- Restart Ollama: ollama serve\n"
                    f"- Re-pull the model: ollama pull {self.model_name}\n"
                    f"Original error: {error_str}"
                )
            print(f"❌ Error generating response: {e}")
            import traceback
            traceback.print_exc()
            return f"Error generating response: {str(e)}"
    
    async def generate_response_stream(
        self,
        message: str,
        context_length: int,
        memory_enabled: bool,
        temperature: float,
        system_prompt: str,
        history: List[Dict],
        top_k: int = 5,
        filter_doc_ids: Optional[List[str]] = None
    ):
        """Generate response with streaming support - yields chunks as they arrive"""
        try:
            if not self.is_available():
                yield "data: " + json.dumps({"chunk": "⚠️ Ollama is not running. Please start it with:\n\n1. Open a terminal\n2. Run: ollama serve\n3. Ensure model is pulled: ollama pull gemma2:2b", "done": True}) + "\n\n"
                return
            
            query = message.strip()
            if not query:
                yield "data: " + json.dumps({"chunk": "Please provide a question or query.", "done": True}) + "\n\n"
                return
            
            # Determine if this needs document context
            needs_context = self._needs_document_context(query, filter_doc_ids)
            context = ""
            citations = []
            
            if needs_context:
                # Retrieve context with citations (async)
                query_lower = query.lower()
                is_analysis_query = any(word in query_lower for word in ["summar", "conclusion", "analyze", "synthesize", "overview", "key points", "main ideas", "summary"])
                retrieval_k = top_k * 2 if is_analysis_query else top_k
                
                print(f"🔍 Retrieving context for query: '{query[:100]}...'")
                if filter_doc_ids:
                    print(f"📄 Filtering to selected documents: {filter_doc_ids}")
                else:
                    print(f"📚 Searching all documents in knowledge base")
                
                context, citations = await self.retrieve_context(query, top_k=retrieval_k, min_similarity=0.2, filter_doc_ids=filter_doc_ids)
                
                if context:
                    print(f"✅ Retrieved {len(citations)} relevant chunks ({len(context)} chars)")
                    print(f"   Context will be pipelined to LLM")
                else:
                    print("⚠️ No relevant context found - LLM will answer without document context")
            else:
                print(f"💬 Conversational message detected - responding naturally without document search")
            
            # Build conversation history
            conversation = []
            if memory_enabled and len(history) > 1:
                max_history = min(len(history) - 1, context_length // 512)
                for msg in history[-max_history:]:
                    role = msg.get('role', 'user')
                    content = msg.get('content', '')
                    if role == 'user':
                        conversation.append(f"User: {content}")
                    elif role == 'assistant':
                        conversation.append(f"Assistant: {content}")
            
            # Build prompt
            prompt = f"{system_prompt}\n\n"
            
            if context and needs_context:
                print(f"📝 Adding context to LLM prompt ({len(context)} characters)")
                prompt += f"=== RELEVANT DOCUMENT CONTEXT ===\n"
                prompt += f"The following information has been retrieved from {len(citations)} document chunks from the user's documents.\n\n"
                prompt += f"{context}\n\n"
                prompt += f"=== ANSWERING INSTRUCTIONS ===\n"
                prompt += f"You are an intelligent assistant. Follow these steps to answer the question:\n\n"
                prompt += f"**STEP 1: Try to answer from context FIRST**\n"
                prompt += f"1. Carefully read the document context provided above.\n"
                prompt += f"2. Check if the context contains information that directly answers the question.\n"
                prompt += f"3. If the context has sufficient information:\n"
                prompt += f"   - Answer based PRIMARILY on the context provided\n"
                prompt += f"   - Cite sources using [1], [2], etc. when referencing specific information\n"
                prompt += f"   - Analyze and synthesize information from multiple chunks if needed\n"
                prompt += f"   - Draw logical conclusions from the provided context\n\n"
                prompt += f"**STEP 2: If context is insufficient or doesn't answer the question**\n"
                prompt += f"4. If the context doesn't contain enough information to fully answer the question:\n"
                prompt += f"   - First, state what information IS available in the context (if any)\n"
                prompt += f"   - Then, clearly indicate that the answer is not fully available in the provided documents\n"
                prompt += f"   - Use this format: \"The provided documents don't contain [specific information], but I can tell you from my knowledge that...\"\n"
                prompt += f"   - Provide your answer using your general knowledge\n"
                prompt += f"   - Be helpful and comprehensive in your response\n\n"
                prompt += f"**IMPORTANT GUIDELINES:**\n"
                prompt += f"- Always prioritize context when it's available and relevant\n"
                prompt += f"- When using general knowledge, always acknowledge it clearly\n"
                prompt += f"- If context partially answers the question, use it and supplement with general knowledge if needed\n"
                prompt += f"- Be transparent about what comes from documents vs. your knowledge\n\n"
            elif needs_context and not context:
                # Document search was attempted but no context found
                prompt += f"=== ANSWERING INSTRUCTIONS ===\n"
                prompt += f"No relevant documents were found in the knowledge base for this question.\n\n"
                prompt += f"You should:\n"
                prompt += f"1. Clearly state: \"The provided documents don't contain information about this topic.\"\n"
                prompt += f"2. Then provide a helpful answer using your general knowledge\n"
                prompt += f"3. Use this format: \"However, I can tell you from my knowledge that...\"\n"
                prompt += f"4. Be comprehensive and helpful in your response\n\n"
            else:
                # Conversational message - respond naturally
                prompt += f"=== INSTRUCTIONS ===\n"
                prompt += f"This is a conversational message. Respond naturally and helpfully without mentioning documents.\n"
                prompt += f"Just have a friendly conversation - no need to reference documents or context.\n\n"
            
            if conversation:
                prompt += f"=== CONVERSATION HISTORY ===\n" + '\n'.join(conversation) + "\n\n"
            
            prompt += f"=== QUESTION ===\n{query}\n\n=== ANSWER ===\n"
            
            # Calculate required context window size
            # Estimate tokens: ~4 characters per token, add buffer for safety
            prompt_tokens = len(prompt) // 4
            required_ctx = max(prompt_tokens + 1000, context_length)  # Ensure enough space for response
            
            print(f"📏 Prompt size: {len(prompt)} chars (~{prompt_tokens} tokens), setting context window to {required_ctx}")
            if context:
                print(f"🔍 Prompt preview (first 500 chars): {prompt[:500]}...")
                print(f"🔍 Prompt contains context: {'RELEVANT DOCUMENT CONTEXT' in prompt}")
            
            # Call Ollama with streaming - use queue for real-time streaming
            import asyncio
            import threading
            queue = asyncio.Queue()
            done_flag = asyncio.Event()
            loop = asyncio.get_event_loop()
            
            def stream_ollama():
                try:
                    response = requests.post(
                        f"{self.ollama_url}/api/generate",
                        json={
                            "model": self.model_name,
                            "prompt": prompt,
                            "stream": True,
                            "options": {
                                "temperature": temperature,
                                "num_ctx": required_ctx,  # Context window size - CRITICAL for RAG!
                                "num_predict": min(context_length, 2048)  # Max tokens to generate
                            }
                        },
                        stream=True,
                        timeout=120
                    )
                    
                    if response.status_code == 200:
                        for line in response.iter_lines(decode_unicode=True):
                            if line:
                                try:
                                    data = json.loads(line)
                                    chunk = data.get('response', '')
                                    if chunk:
                                        # Put chunk in queue immediately
                                        asyncio.run_coroutine_threadsafe(
                                            queue.put(("data: " + json.dumps({"chunk": chunk, "done": False}) + "\n\n")),
                                            loop
                                        )
                                    
                                    if data.get('done', False):
                                        if citations:
                                            citations_text = "\n\n=== Sources ===\n"
                                            for citation in citations:
                                                citations_text += f"{citation['number']}. {citation['source']}"
                                                if citation.get('page'):
                                                    citations_text += f", page {citation['page']}"
                                                citations_text += f" (similarity: {citation['similarity']:.3f})\n"
                                            asyncio.run_coroutine_threadsafe(
                                                queue.put(("data: " + json.dumps({"chunk": citations_text, "done": True}) + "\n\n")),
                                                loop
                                            )
                                        else:
                                            asyncio.run_coroutine_threadsafe(
                                                queue.put(("data: " + json.dumps({"chunk": "", "done": True}) + "\n\n")),
                                                loop
                                            )
                                        loop.call_soon_threadsafe(done_flag.set)
                                        break
                                except json.JSONDecodeError:
                                    continue
                    else:
                        error_text = response.text
                        # Check for model loading errors
                        if "load" in error_text.lower() and ("EOF" in error_text or "connection" in error_text.lower()):
                            error_msg = (
                                f"Error: Model loading failed. This usually means:\n"
                                f"1. The model '{self.model_name}' is not available or corrupted\n"
                                f"2. Ollama is having trouble loading the model\n"
                                f"3. There's insufficient memory/disk space\n\n"
                                f"Try:\n"
                                f"- Check if model exists: ollama list\n"
                                f"- Restart Ollama: ollama serve\n"
                                f"- Re-pull the model: ollama pull {self.model_name}\n"
                                f"Original error: {response.status_code} - {error_text}"
                            )
                        else:
                            error_msg = f"Error: {response.status_code} - {error_text}"
                        error_chunk = "data: " + json.dumps({"chunk": error_msg, "done": True}) + "\n\n"
                        asyncio.run_coroutine_threadsafe(queue.put(error_chunk), loop)
                        loop.call_soon_threadsafe(done_flag.set)
                except requests.exceptions.ConnectionError as e:
                    error_msg = (
                        f"❌ Cannot connect to Ollama.\n\n"
                        f"Please ensure Ollama is running: ollama serve\n"
                        f"Error: {str(e)}"
                    )
                    error_chunk = "data: " + json.dumps({"chunk": error_msg, "done": True}) + "\n\n"
                    asyncio.run_coroutine_threadsafe(queue.put(error_chunk), loop)
                    loop.call_soon_threadsafe(done_flag.set)
                except Exception as e:
                    error_str = str(e)
                    # Check for EOF/model loading errors in exception message
                    if "load" in error_str.lower() and ("EOF" in error_str or "connection" in error_str.lower()):
                        error_msg = (
                            f"Error: Model loading failed. This usually means:\n"
                            f"1. The model '{self.model_name}' is not available or corrupted\n"
                            f"2. Ollama is having trouble loading the model\n"
                            f"3. There's insufficient memory/disk space\n\n"
                            f"Try:\n"
                            f"- Check if model exists: ollama list\n"
                            f"- Restart Ollama: ollama serve\n"
                            f"- Re-pull the model: ollama pull {self.model_name}\n"
                            f"Original error: {error_str}"
                        )
                    else:
                        error_msg = f"Error: {str(e)}"
                    error_chunk = "data: " + json.dumps({"chunk": error_msg, "done": True}) + "\n\n"
                    asyncio.run_coroutine_threadsafe(queue.put(error_chunk), loop)
                    loop.call_soon_threadsafe(done_flag.set)
            
            # Start streaming in thread
            thread = threading.Thread(target=stream_ollama, daemon=True)
            thread.start()
            
            # Yield chunks as they arrive
            while not done_flag.is_set() or not queue.empty():
                try:
                    chunk = await asyncio.wait_for(queue.get(), timeout=0.1)
                    yield chunk
                    if '"done":true' in chunk:
                        break
                except asyncio.TimeoutError:
                    if done_flag.is_set():
                        break
                    continue
            
        except Exception as e:
            print(f"❌ Error generating streaming response: {e}")
            import traceback
            traceback.print_exc()
            yield "data: " + json.dumps({"chunk": f"Error generating response: {str(e)}", "done": True}) + "\n\n"
    
    def is_available(self) -> bool:
        """Check if Ollama is available"""
        try:
            response = requests.get(f"{self.ollama_url}/api/tags", timeout=5)
            return response.status_code == 200
        except:
            return False
    
    def _preload_model(self) -> bool:
        """Preload the model to avoid EOF errors during first request"""
        try:
            print(f"🔄 Preloading model {self.model_name}...")
            # Make a small test request to load the model
            response = requests.post(
                f"{self.ollama_url}/api/generate",
                json={
                    "model": self.model_name,
                    "prompt": "test",
                    "stream": False,
                    "options": {
                        "num_predict": 1,  # Just generate 1 token to load the model
                        "num_ctx": 128
                    }
                },
                timeout=60
            )
            if response.status_code == 200:
                print(f"✅ Model {self.model_name} preloaded successfully")
                return True
            else:
                print(f"⚠️ Model preload returned status {response.status_code}: {response.text}")
                return False
        except Exception as e:
            print(f"⚠️ Model preload failed: {e}")
            return False
    
    def switch_model(self, model_name: str) -> bool:
        """Switch to a different Ollama model and preload it"""
        try:
            # Verify model exists
            response = requests.get(f"{self.ollama_url}/api/tags", timeout=5)
            if response.status_code == 200:
                models = response.json().get('models', [])
                model_names = [m.get('name', '') for m in models]
                if model_name in model_names:
                    old_model = self.model_name
                    self.model_name = model_name
                    print(f"🔄 Switching RAG model from {old_model} to {model_name}...")
                    
                    # Give Ollama a moment to unload the previous model
                    import time
                    time.sleep(0.5)
                    
                    # Preload the new model to avoid EOF errors on first request
                    if self._preload_model():
                        print(f"✅ Switched RAG model to: {model_name} (preloaded)")
                        return True
                    else:
                        print(f"⚠️ Switched RAG model to: {model_name} (preload failed, but will try on first request)")
                        # Still return True - the model switch succeeded, preload is just optimization
                        return True
                else:
                    print(f"⚠️ Model {model_name} not found. Available models: {model_names}")
                    return False
            return False
        except Exception as e:
            print(f"❌ Error switching model: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def get_available_models(self) -> List[str]:
        """Get list of available Ollama models"""
        try:
            response = requests.get(f"{self.ollama_url}/api/tags", timeout=5)
            if response.status_code == 200:
                models = response.json().get('models', [])
                return [m.get('name', '') for m in models]
            return []
        except:
            return []
    
    def download_model(self, model_name: str) -> Dict[str, Any]:
        """Download an Ollama model"""
        try:
            response = requests.post(
                f"{self.ollama_url}/api/pull",
                json={"name": model_name},
                stream=True,
                timeout=300
            )
            
            if response.status_code == 200:
                # Stream the download progress
                progress_data = []
                for line in response.iter_lines():
                    if line:
                        try:
                            data = json.loads(line)
                            progress_data.append(data)
                            if data.get('status') == 'success':
                                return {
                                    'success': True,
                                    'message': f'Model {model_name} downloaded successfully',
                                    'progress': progress_data
                                }
                        except:
                            pass
                
                return {
                    'success': True,
                    'message': f'Model {model_name} download initiated',
                    'progress': progress_data
                }
            else:
                return {
                    'success': False,
                    'error': f'Failed to download model: {response.status_code}'
                }
        except Exception as e:
            return {
                'success': False,
                'error': f'Error downloading model: {str(e)}'
            }
